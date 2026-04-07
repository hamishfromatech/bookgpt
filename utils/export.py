"""
Export utilities for BookGPT application.

Supports multiple export formats: PDF, EPUB, DOCX, and plain text.
"""

import os
import io
import json
import logging
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("reportlab not installed. PDF export will be limited.")

try:
    from ebooklib import epub
    from ebooklib.epub import EpubBook, EpubHtml
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False
    logger.warning("ebooklib not installed. EPUB export will be unavailable.")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX export will be unavailable.")


class BookExporter:
    """Handles exporting books to various formats."""

    def __init__(self, project_id: str, project_dir: str = "projects"):
        self.project_id = project_id
        self.project_dir = os.path.join(project_dir, project_id)

    def _read_chapters(self) -> List[Dict[str, Any]]:
        """Read all chapter files for the project."""
        chapters = []
        chapters_dir = os.path.join(self.project_dir, "chapters")

        if not os.path.exists(chapters_dir):
            return chapters

        # Get all chapter files and sort them naturally
        chapter_files = []
        for filename in os.listdir(chapters_dir):
            if filename.endswith('.md'):
                try:
                    # Extract chapter number from filename
                    num = int(filename.split('_')[1].split('.')[0])
                    chapter_files.append((num, filename))
                except (IndexError, ValueError):
                    continue

        chapter_files.sort(key=lambda x: x[0])

        for num, filename in chapter_files:
            filepath = os.path.join(chapters_dir, filename)
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                chapters.append({
                    'number': num,
                    'filename': filename,
                    'content': content,
                    'word_count': len(content.split())
                })
            except Exception as e:
                logger.error(f"Error reading chapter {filename}: {e}")

        return chapters

    def _read_project_files(self) -> Dict[str, str]:
        """Read project outline and research notes."""
        files = {}

        for filename in ['outline.md', 'research_notes.md', 'editing_notes.md']:
            filepath = os.path.join(self.project_dir, filename)
            if os.path.exists(filepath):
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        files[filename] = f.read()
                except Exception as e:
                    logger.warning(f"Could not read {filename}: {e}")

        return files

    def export_to_text(self) -> bytes:
        """Export book as plain text."""
        chapters = self._read_chapters()
        project_files = self._read_project_files()

        lines = []

        # Add outline if available
        if 'outline.md' in project_files:
            lines.append("=" * 60)
            lines.append("BOOK OUTLINE")
            lines.append("=" * 60)
            lines.append(project_files['outline.md'])
            lines.append("\n\n")

        # Add chapters
        lines.append("=" * 60)
        lines.append("BOOK CONTENT")
        lines.append("=" * 60)
        lines.append("\n")

        for chapter in chapters:
            lines.append("-" * 40)
            lines.append(f"CHAPTER {chapter['number']}")
            lines.append("-" * 40)
            lines.append("")
            lines.append(chapter['content'])
            lines.append("\n")

        content = "\n".join(lines)
        return content.encode('utf-8')

    def export_to_json(self) -> bytes:
        """Export book as JSON with full metadata."""
        chapters = self._read_chapters()
        project_files = self._read_project_files()

        data = {
            'project_id': self.project_id,
            'exported_at': datetime.now().isoformat(),
            'chapters': chapters,
            'outline': project_files.get('outline.md', ''),
            'research_notes': project_files.get('research_notes.md', ''),
            'editing_notes': project_files.get('editing_notes.md', ''),
            'metadata': {
                'total_chapters': len(chapters),
                'total_words': sum(c['word_count'] for c in chapters)
            }
        }

        return json.dumps(data, indent=2).encode('utf-8')

    def export_to_pdf(self, title: str = "Book", author: str = "Unknown",
                      page_size: str = "letter") -> Optional[bytes]:
        """Export book as PDF."""
        if not REPORTLAB_AVAILABLE:
            logger.error("reportlab not installed. Cannot export to PDF.")
            return None

        chapters = self._read_chapters()

        # Create buffer
        buffer = io.BytesIO()

        # Select page size
        pagesize = letter if page_size == "letter" else A4

        # Create document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=pagesize,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        # Styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='BookTitle',
            parent=styles['Heading1'],
            fontSize=24,
            alignment=TA_CENTER,
            spaceAfter=30
        ))
        styles.add(ParagraphStyle(
            name='ChapterTitle',
            parent=styles['Heading2'],
            fontSize=18,
            spaceAfter=12,
            spaceBefore=24
        ))
        styles.add(ParagraphStyle(
            name='BodyText',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            leading=14
        ))

        # Build content
        story = []

        # Title page
        story.append(Spacer(1, 2*inch))
        story.append(Paragraph(title, styles['BookTitle']))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(f"By {author}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(PageBreak())

        # Chapters
        for chapter in chapters:
            story.append(Paragraph(f"Chapter {chapter['number']}", styles['ChapterTitle']))

            # Split content into paragraphs
            paragraphs = chapter['content'].split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if para:
                    # Escape special characters for reportlab
                    para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    try:
                        story.append(Paragraph(para, styles['BodyText']))
                    except Exception as e:
                        logger.warning(f"Could not render paragraph: {e}")
                        story.append(Paragraph(para[:200] + "...", styles['BodyText']))

            story.append(PageBreak())

        # Build PDF
        doc.build(story)

        buffer.seek(0)
        return buffer.getvalue()

    def export_to_epub(self, title: str = "Book", author: str = "Unknown",
                       language: str = "en") -> Optional[bytes]:
        """Export book as EPUB."""
        if not EPUB_AVAILABLE:
            logger.error("ebooklib not installed. Cannot export to EPUB.")
            return None

        chapters = self._read_chapters()

        # Create EPUB book
        book = EpubBook()
        book.set_identifier(f"book-{self.project_id}-{datetime.now().timestamp()}")
        book.set_title(title)
        book.set_language(language)
        book.add_author(author)

        # Create chapter items
        chapter_items = []
        toc_links = []

        for chapter in chapters:
            # Create HTML content
            html_content = f"""
            <html xmlns="http://www.w3.org/1999/xhtml">
            <head><title>Chapter {chapter['number']}</title></head>
            <body>
                <h1>Chapter {chapter['number']}</h1>
                {self._markdown_to_html(chapter['content'])}
            </body>
            </html>
            """

            chapter_item = EpubHtml(
                uid=f"chapter_{chapter['number']}",
                file_name=f"chapter_{chapter['number']}.xhtml",
                content=html_content.encode('utf-8')
            )
            book.add_item(chapter_item)
            chapter_items.append(chapter_item)
            toc_links.append((f"Chapter {chapter['number']}", chapter_item))

        # Add TOC
        book.toc = tuple(toc_links)

        # Add navigation files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Set spine
        book.spine = ['nav'] + chapter_items

        # Write to buffer
        buffer = io.BytesIO()
        epub.write_epub(buffer, book, {})
        buffer.seek(0)

        return buffer.getvalue()

    def export_to_docx(self, title: str = "Book", author: str = "Unknown") -> Optional[bytes]:
        """Export book as DOCX."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed. Cannot export to DOCX.")
            return None

        chapters = self._read_chapters()

        # Create document
        doc = Document()

        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Author
        author_para = doc.add_paragraph(f"By {author}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Generation date
        date_para = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Chapters
        for chapter in chapters:
            doc.add_heading(f"Chapter {chapter['number']}", level=1)

            # Split into paragraphs
            paragraphs = chapter['content'].split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if para:
                    doc.add_paragraph(para)

            doc.add_page_break()

        # Write to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def _markdown_to_html(self, markdown_text: str) -> str:
        """Convert basic markdown to HTML for EPUB."""
        # Simple conversion - in production use a proper markdown library
        html = markdown_text

        # Convert headers
        html = html.replace('\n# ', '\n<h1>').replace('\n## ', '\n<h2>')

        # Convert paragraphs
        paragraphs = html.split('\n\n')
        html_paragraphs = []
        for para in paragraphs:
            para = para.strip()
            if para:
                if not para.startswith('<h'):
                    html_paragraphs.append(f'<p>{para}</p>')
                else:
                    html_paragraphs.append(para)

        return '\n'.join(html_paragraphs)

    def export(self, format: str, title: str = "Book", author: str = "Unknown") -> Optional[bytes]:
        """
        Export book to specified format.

        Args:
            format: One of 'txt', 'json', 'pdf', 'epub', 'docx'
            title: Book title
            author: Author name

        Returns:
            bytes: The exported file content, or None if export failed
        """
        format = format.lower()

        if format == 'txt' or format == 'text':
            return self.export_to_text()
        elif format == 'json':
            return self.export_to_json()
        elif format == 'pdf':
            return self.export_to_pdf(title, author)
        elif format == 'epub':
            return self.export_to_epub(title, author)
        elif format == 'docx' or format == 'doc':
            return self.export_to_docx(title, author)
        else:
            logger.error(f"Unknown export format: {format}")
            return None

    def get_available_formats(self) -> List[str]:
        """Get list of available export formats."""
        formats = ['txt', 'json']

        if REPORTLAB_AVAILABLE:
            formats.append('pdf')
        if EPUB_AVAILABLE:
            formats.append('epub')
        if DOCX_AVAILABLE:
            formats.append('docx')

        return formats